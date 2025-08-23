"""
Document processor for converting various file formats to plain text.
Supports PDF, Word, Excel, and text files with optional OCR capabilities.
"""

import os
import re
from pathlib import Path
from typing import List, Optional
import logging

# Document processing libraries
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# OCR support
try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Main document processor class for handling various file formats"""
    
    def __init__(self, chunk_size: int = 1000, enable_ocr: bool = False):
        """
        Initialize the document processor
        
        Args:
            chunk_size: Maximum size of text chunks
            enable_ocr: Whether to enable OCR for scanned PDFs
        """
        self.chunk_size = chunk_size
        self.enable_ocr = enable_ocr and OCR_AVAILABLE
        
        if not OCR_AVAILABLE and enable_ocr:
            logger.warning("OCR requested but pytesseract/pdf2image not available")
    
    def process_document(self, file_path: str) -> List[str]:
        """
        Process a document and return text chunks
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of text chunks
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file type and process accordingly
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            text = self._process_pdf(file_path)
        elif file_extension == '.docx':
            text = self._process_docx(file_path)
        elif file_extension in ['.xlsx', '.xls']:
            text = self._process_excel(file_path)
        elif file_extension == '.txt':
            text = self._process_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Clean and chunk the text
        cleaned_text = self._clean_text(text)
        chunks = self._chunk_text(cleaned_text)
        
        logger.info(f"Processed {file_path.name}: {len(chunks)} chunks created")
        return chunks
    
    def _process_pdf(self, file_path: Path) -> str:
        """Process PDF files with fallback to OCR if needed"""
        text = ""
        
        # Try pdfplumber first (better for text-based PDFs)
        if PDFPLUMBER_AVAILABLE:
            try:
                text = self._extract_text_with_pdfplumber(file_path)
                if text.strip():
                    logger.info("Successfully extracted text using pdfplumber")
                    return text
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}")
        
        # Try PyMuPDF as fallback
        if PYMUPDF_AVAILABLE:
            try:
                text = self._extract_text_with_pymupdf(file_path)
                if text.strip():
                    logger.info("Successfully extracted text using PyMuPDF")
                    return text
            except Exception as e:
                logger.warning(f"PyMuPDF failed: {e}")
        
        # Try OCR if enabled and other methods failed
        if self.enable_ocr and OCR_AVAILABLE:
            try:
                text = self._extract_text_with_ocr(file_path)
                if text.strip():
                    logger.info("Successfully extracted text using OCR")
                    return text
            except Exception as e:
                logger.warning(f"OCR failed: {e}")
        
        # If all methods failed, raise an error
        if not text.strip():
            raise RuntimeError("Failed to extract text from PDF using all available methods")
        
        return text
    
    def _extract_text_with_pdfplumber(self, file_path: Path) -> str:
        """Extract text using pdfplumber"""
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    
    def _extract_text_with_pymupdf(self, file_path: Path) -> str:
        """Extract text using PyMuPDF"""
        text = ""
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    
    def _extract_text_with_ocr(self, file_path: Path) -> str:
        """Extract text using OCR (Tesseract)"""
        text = ""
        
        # Convert PDF pages to images
        images = convert_from_path(file_path, dpi=300)
        
        # Extract text from each image
        for i, image in enumerate(images):
            logger.info(f"Processing page {i+1}/{len(images)} with OCR")
            page_text = pytesseract.image_to_string(image, lang='eng')
            text += page_text + "\n"
        
        return text
    
    def _process_docx(self, file_path: Path) -> str:
        """Process Word documents"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word document processing")
        
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text
    
    def _process_excel(self, file_path: Path) -> str:
        """Process Excel files"""
        text = ""
        
        # Try pandas first
        if PANDAS_AVAILABLE:
            try:
                # Read all sheets
                excel_file = pd.ExcelFile(file_path)
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    text += f"Sheet: {sheet_name}\n"
                    text += df.to_string(index=False) + "\n\n"
                return text
            except Exception as e:
                logger.warning(f"pandas failed: {e}")
        
        # Fallback to openpyxl
        if OPENPYXL_AVAILABLE:
            try:
                workbook = openpyxl.load_workbook(file_path, data_only=True)
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text += f"Sheet: {sheet_name}\n"
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                        text += row_text + "\n"
                    text += "\n"
                
                return text
            except Exception as e:
                logger.warning(f"openpyxl failed: {e}")
        
        raise RuntimeError("Failed to process Excel file with available libraries")
    
    def _process_txt(self, file_path: Path) -> str:
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            raise RuntimeError("Failed to decode text file with any supported encoding")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere with LLM processing
        text = re.sub(r'[^\w\s\-.,;:!?()\[\]{}"\']', ' ', text)
        
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        
        # Remove empty lines
        text = '\n'.join(line.strip() for line in text.split('\n') if line.strip())
        
        return text.strip()
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into manageable chunks"""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        current_chunk = ""
        
        # Split by sentences first
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        for sentence in sentences:
            # If adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = sentence
            else:
                current_chunk += " " + sentence if current_chunk else sentence
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        # If chunks are still too long, split by words
        final_chunks = []
        for chunk in chunks:
            if len(chunk) <= self.chunk_size:
                final_chunks.append(chunk)
            else:
                # Split by words
                words = chunk.split()
                current_word_chunk = ""
                
                for word in words:
                    if len(current_word_chunk) + len(word) + 1 > self.chunk_size and current_word_chunk:
                        final_chunks.append(current_word_chunk.strip())
                        current_word_chunk = word
                    else:
                        current_word_chunk += " " + word if current_word_chunk else word
                
                if current_word_chunk.strip():
                    final_chunks.append(current_word_chunk.strip())
        
        return final_chunks
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        formats = []
        
        if PDFPLUMBER_AVAILABLE or PYMUPDF_AVAILABLE:
            formats.append("PDF (.pdf)")
        
        if DOCX_AVAILABLE:
            formats.append("Word (.docx)")
        
        if PANDAS_AVAILABLE or OPENPYXL_AVAILABLE:
            formats.append("Excel (.xlsx, .xls)")
        
        formats.append("Text (.txt)")
        
        return formats 